import os
import base64
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from PIL import Image

default_styles = {
    'body': 'font-family: system-ui, -apple-system, sans-serif; max-width: 800px; margin: 20px auto; padding: 20px; background: #f8fafc; color: #333;',
    'wrapper': 'background: white; padding: 30px; border-radius: 10px;',
    'h1': 'color: #1a365d; font-size: 32px; margin: 15px 0; padding-bottom: 10px; border-bottom: 2px solid #3498db;',
    'h2': 'color: #2c5282; font-size: 24px; margin: 15px 0; padding-left: 15px;',
    'h3': 'color: #2b6cb0; font-size: 20px; margin: 15px 0;',
    'h4': 'color: #2b6cb0; font-size: 19px; margin: 15px 0;',
    'h5': 'color: #2b6cb0; font-size: 18px; margin: 15px 0;',
    'h6': 'color: #2b6cb0; font-size: 17px; margin: 15px 0;',
    'p': 'line-height: 1.6; margin: 15px 0; color: #2d3748;',
    'image': 'margin: 20px auto; text-align: center;',
    'img': 'max-width: 100%; height: auto; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);',
    'strong': 'color: #2c5282; background: #e2e8f0; padding: 2px 5px; border-radius: 3px;',
    'em': 'color: #4a5568; font-style: italic;',
}


def _get_image_mime_type(image_data):
    magic_numbers = {
        b'\x89PNG\r\n\x1a\n': 'image/png',
        b'\xFF\xD8\xFF': 'image/jpeg',
        b'GIF87a': 'image/gif',
        b'GIF89a': 'image/gif',
        b'\x00\x00\x01\x00': 'image/x-icon',
    }
    for magic, mime in magic_numbers.items():
        if image_data.startswith(magic):
            return mime
    return 'image/png'


def _convert_image_to_base64(image_data):
    try:
        img = Image.open(BytesIO(image_data))
        if img.mode in ('RGBA', 'LA'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[-1])
            img = background

        max_width = 700
        if img.size[0] > max_width:
            ratio = max_width / img.size[0]
            img = img.resize((max_width, int(img.size[1] * ratio)), Image.Resampling.LANCZOS)

        output = BytesIO()
        img.save(output, format='JPEG', quality=85, optimize=True)
        image_data = output.getvalue()
        mime_type = 'image/jpeg'
    except Exception as e:
        print(f"Error procesando imagen: {e}")
        mime_type = _get_image_mime_type(image_data)

    base64_data = base64.b64encode(image_data).decode('utf-8')
    return f'data:{mime_type};base64,{base64_data}'


def find_image_id(element):
    blip = element.find('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
    if blip is not None:
        embed = blip.get(qn('r:embed'))
        return embed
    return None


def is_list_paragraph(para):
    """Detecta si el párrafo es parte de una lista (viñetas o numerada)"""
    p = para._p
    numPr = p.find('w:pPr/w:numPr', namespaces=p.nsmap)
    return numPr is not None


def get_list_type(para):
    """Devuelve 'ul' o 'ol' según si es lista con viñetas o numerada"""
    try:
        if "number" in para.style.name.lower():
            return "ol"
        return "ul"
    except:
        return "ul"


def build_html(file_path, output_path=None, styles=None):
    doc = Document(file_path)
    html_content = []

    if not styles:
        styles = default_styles

    html_content.extend([
        '<!DOCTYPE html>',
        '<html lang="es">',
        '<head>',
        '<meta charset="UTF-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        '<title>Documento Convertido</title>',
        '</head>',
        f'<body style="{styles["body"]}">',
        f'<div style="{styles["wrapper"]}">'
    ])

    try:
        # Mapear imágenes
        image_rels = {}
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    image_base64 = _convert_image_to_base64(image_data)
                    image_rels[rel.rId] = image_base64
                except Exception as e:
                    print(f"Error cargando imagen: {e}")

        # Procesar párrafos
        in_list = False
        list_tag = None

        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = getattr(para.style, "name", "").lower()

            # Encabezados
            if style_name.startswith("heading"):
                try:
                    level = int(style_name.split()[-1])
                    heading_style = styles.get(f"h{level}", "")
                    html_content.append(f'<h{level} style="{heading_style}">{text}</h{level}>')
                    continue
                except (IndexError, ValueError):
                    pass

            # Verificar si es lista
            if is_list_paragraph(para):
                tag = get_list_type(para)

                # Abrir lista si necesario
                if not in_list or list_tag != tag:
                    if in_list:
                        html_content.append(f'</{list_tag}>')
                    html_content.append(f'<{tag}>')
                    in_list = True
                    list_tag = tag

                # Agregar item de lista
                formatted_text = ""
                for run in para.runs:
                    run_text = run.text
                    if not run_text.strip():
                        continue
                    if run.bold and run.italic:
                        formatted_text += f'<strong style="{styles["strong"]}"><em style="{styles["em"]}">{run_text}</em></strong>'
                    elif run.bold:
                        formatted_text += f'<strong style="{styles["strong"]}">{run_text}</strong>'
                    elif run.italic:
                        formatted_text += f'<em style="{styles["em"]}">{run_text}</em>'
                    else:
                        formatted_text += run_text

                html_content.append(f'<li>{formatted_text}</li>')
                continue
            else:
                if in_list:
                    html_content.append(f'</{list_tag}>')
                    in_list = False
                    list_tag = None

            # Imágenes en línea
            has_images = False
            xml_element = para._element
            for element in xml_element.iter():
                if element.tag.endswith('drawing'):
                    rel_id = find_image_id(element)
                    if rel_id and rel_id in image_rels:
                        html_content.append(f'<div style="{styles["image"]}">')
                        html_content.append(f'<img src="{image_rels[rel_id]}" alt="Imagen del documento" loading="lazy" style="{styles["img"]}">')
                        html_content.append('</div>')
                        has_images = True

            # Texto normal
            if text or not has_images:
                formatted_text = ""
                for run in para.runs:
                    run_text = run.text
                    if not run_text.strip():
                        continue
                    if run.bold and run.italic:
                        formatted_text += f'<strong style="{styles["strong"]}"><em style="{styles["em"]}">{run_text}</em></strong>'
                    elif run.bold:
                        formatted_text += f'<strong style="{styles["strong"]}">{run_text}</strong>'
                    elif run.italic:
                        formatted_text += f'<em style="{styles["em"]}">{run_text}</em>'
                    else:
                        formatted_text += run_text

                if formatted_text.strip():
                    html_content.append(f'<p style="{styles["p"]}">{formatted_text}</p>')

        # Cerrar lista si quedó abierta
        if in_list:
            html_content.append(f'</{list_tag}>')

        # Procesar tablas
        for table in doc.tables:
            html_content.append('<table border="1" style="border-collapse: collapse; margin: 20px 0; width: 100%;">')
            for row in table.rows:
                html_content.append('<tr>')
                for cell in row.cells:
                    cell_text = ' '.join(paragraph.text.strip() for paragraph in cell.paragraphs)
                    html_content.append(f'<td style="padding: 8px; border: 1px solid #ccc;">{cell_text}</td>')
                html_content.append('</tr>')
            html_content.append('</table>')

    except Exception as e:
        print(f"Error procesando documento: {e}")

    html_content.append('</div>')
    html_content.append('</body>')
    html_content.append('</html>')

    final_html = '\n'.join(html_content)
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_html)

    return final_html
